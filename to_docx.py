import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import os

def set_cell_background(cell, fill_hex):
    """Sets the background color of a Word table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets margins (padding) of a cell in twentieths of a point (dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def parse_markdown_to_docx(md_path, docx_path):
    print(f"Reading Markdown file: {md_path}...")
    if not os.path.exists(md_path):
        raise FileNotFoundError(f"File {md_path} not found.")

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style definitions
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(44, 62, 80) # Dark Charcoal/Blue

    # Regular expressions for inline markdown formatting
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')
    italic_pattern = re.compile(r'\*(.*?)\*')
    link_pattern = re.compile(r'\[(.*?)\]\((.*?)\)')
    image_pattern = re.compile(r'\!\[(.*?)\]\((.*?)\)')

    def format_paragraph_with_runs(paragraph, text):
        # Temp placeholders for styling
        # Find bold and italic tokens
        parts = []
        last_idx = 0
        
        # We parse images separately, skip if it's an image block
        if text.startswith('!['):
            return
            
        # Simplistic parser for bold and italic
        tokens = []
        for m in bold_pattern.finditer(text):
            tokens.append((m.start(), m.end(), 'bold', m.group(1)))
        for m in italic_pattern.finditer(text):
            # Ensure it's not part of a bold pattern
            is_bold = False
            for b_start, b_end, _, _ in tokens:
                if b_start <= m.start() < b_end:
                    is_bold = True
                    break
            if not is_bold:
                tokens.append((m.start(), m.end(), 'italic', m.group(1)))
                
        tokens.sort(key=lambda x: x[0])
        
        current_pos = 0
        for start, end, style, val in tokens:
            if start > current_pos:
                paragraph.add_run(text[current_pos:start])
            run = paragraph.add_run(val)
            if style == 'bold':
                run.bold = True
            elif style == 'italic':
                run.italic = True
            current_pos = end
            
        if current_pos < len(text):
            paragraph.add_run(text[current_pos:])

    # Parser state
    in_table = False
    table_headers = []
    table_rows = []

    def flush_table():
        nonlocal in_table, table_headers, table_rows
        if not table_rows:
            in_table = False
            return
        
        print(f"Rendering table with {len(table_rows)} rows...")
        cols_count = len(table_headers)
        table = doc.add_table(rows=1, cols=cols_count)
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        # Style Header Row
        hdr_cells = table.rows[0].cells
        for col_idx, col_name in enumerate(table_headers):
            hdr_cells[col_idx].text = col_name.strip()
            # Set background and padding
            set_cell_background(hdr_cells[col_idx], '2980B9') # JALA blue header
            set_cell_margins(hdr_cells[col_idx], top=120, bottom=120, left=150, right=150)
            # Center alignment for header text
            hdr_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[col_idx].paragraphs[0].runs[0].font.bold = True
            hdr_cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255) # white text
            hdr_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(10)

        # Style Data Rows
        for row_idx, row_data in enumerate(table_rows):
            row = table.add_row()
            # Alternating shading
            bg_color = 'F2F7F9' if row_idx % 2 == 1 else 'FFFFFF'
            for col_idx, cell_value in enumerate(row_data):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    cell.text = cell_value.strip()
                    set_cell_background(cell, bg_color)
                    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                    # format inline bold/italics in tables
                    text_content = cell_value.strip()
                    cell.paragraphs[0].text = "" # Clear plain text
                    format_paragraph_with_runs(cell.paragraphs[0], text_content)
        
        # Reset table state
        table_headers = []
        table_rows = []
        in_table = False
        doc.add_paragraph() # space after table

    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        
        # Check if line is a table divider (e.g. |---|---|)
        if re.match(r'^\|[\s:-|]*\|$', line):
            idx += 1
            continue
            
        # Check if it is a table line
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            idx += 1
            continue
        elif in_table:
            # We hit a non-table line, flush the table
            flush_table()
            
        # Plain text and formatting
        if not line:
            idx += 1
            continue
            
        # Headers
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(line[2:])
            run.font.name = 'Calibri'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(44, 62, 80)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[3:])
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(41, 128, 185)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[4:])
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(127, 140, 141)
        # Bullet list items
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(0)
            format_paragraph_with_runs(p, line[2:])
        # Number list items
        elif re.match(r'^\d+\.\s+', line):
            # Find the position of the list identifier
            prefix_match = re.match(r'^(\d+\.)\s+', line)
            prefix = prefix_match.group(1)
            content = line[len(prefix)+1:]
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(4)
            format_paragraph_with_runs(p, content)
        # Blockquote / Note boxes
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(127, 140, 141)
        # Embedded images
        elif line.startswith('![') or image_pattern.search(line):
            # Parse image name and filepath
            match = image_pattern.search(line)
            if match:
                img_name = match.group(1)
                img_path = match.group(2)
                # Fix path for Windows if it uses file://
                img_path = img_path.replace('file:///', '').replace('%20', ' ')
                img_path = os.path.abspath(img_path)
                if os.path.exists(img_path):
                    print(f"Embedding image: {img_path}")
                    try:
                        doc.add_picture(img_path, width=Inches(6.0))
                        # Center align picture paragraph
                        p = doc.paragraphs[-1]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_after = Pt(12)
                    except Exception as e:
                        print(f"Error embedding picture: {e}")
                else:
                    print(f"Warning: Image file not found at {img_path}")
        # Paragraphs
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            format_paragraph_with_runs(p, line)
            
        idx += 1

    if in_table:
        flush_table()

    # Save document
    doc.save(docx_path)
    print(f"Word document compiled successfully at: {docx_path}")

if __name__ == "__main__":
    md_file = "Laporan_Data_Storytelling_MDS_Udang.md"
    docx_file = "Laporan_Data_Storytelling_MDS_Udang.docx"
    try:
        parse_markdown_to_docx(md_file, docx_file)
    except Exception as e:
        print(f"Error compilation: {e}")
